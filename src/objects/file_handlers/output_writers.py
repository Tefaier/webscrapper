from __future__ import annotations

from io import TextIOWrapper, BytesIO

from bs4 import PageElement
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from PIL import Image

from objects.types.field_types import FieldTypes
from settings.file_handlers_defaults import IMAGE_WIDTH_CM
from utils.binary_converter import convert_binary
from objects.types.custom_exceptions import UnsupportedArgumentsException


class OutputWriter:
    def write_field(self, data, field_type: FieldTypes) -> None:
        if field_type == FieldTypes.Text:
            if isinstance(data, PageElement):
                self.write_text(data.text)
            elif isinstance(data, str):
                self.write_text(data)
        elif field_type == FieldTypes.Image:
            self.write_image(data, "png")

    def write_text(self, string: str) -> None:
        ...

    def write_image(self, data, format: str) -> None:
        ...

    def close(self) -> None:
        ...


class TxtWriter(OutputWriter):
    def __init__(self, full_path: str):
        self._file: TextIOWrapper = open(full_path, "w", encoding="utf-8")

    def write_text(self, string: str) -> None:
        self._file.write(string + "\n")

    def write_image(self, data, format: str) -> None:
        # Images are not supported for plain text; skip silently
        return

    def close(self) -> None:
        self._file.close()


class HtmlWriter(OutputWriter):
    def __init__(self, full_path: str):
        self._file: TextIOWrapper = open(full_path, "w", encoding="utf-8")
        self._file.write("<html>\n<head></head>\n<body>\n")

    def write_text(self, string: str) -> None:
        self._file.write("<p>" + string + "</p>\n")

    def write_image(self, data, format: str) -> None:
        self._file.write(f'<img src="data:image/{format};base64,{convert_binary(data, "string")}">')

    def close(self) -> None:
        self._file.write("</body>")
        self._file.close()


class DocxWriter(OutputWriter):
    def __init__(self, full_path: str, image_width: float = IMAGE_WIDTH_CM):
        self._full_path = full_path
        self._doc = Document()
        self.image_width = image_width

    def write_text(self, string: str) -> None:
        self._doc.add_paragraph(string)

    def write_image(self, data, format: str) -> None:
        try:
            new_file = BytesIO()
            image = Image.open(BytesIO(convert_binary(data, "PIL")), formats=[format])
            image.save(new_file, format='PNG')
            new_file.seek(0)
            p = self._doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r = p.add_run()
            r.add_picture(new_file, width=Cm(self.image_width))
        except Exception as e:
            return

    def close(self) -> None:
        self._doc.save(self._full_path)


class WriterFactory:
    @staticmethod
    def create(full_path: str, file_type: str) -> OutputWriter:
        if file_type == "txt":
            return TxtWriter(full_path)
        if file_type == "html":
            return HtmlWriter(full_path)
        if file_type == "docx":
            return DocxWriter(full_path)
        raise UnsupportedArgumentsException(f"Unsupported file type: {file_type}")

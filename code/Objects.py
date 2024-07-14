import base64
import io
from enum import Enum

from bs4 import BeautifulSoup
from bs4.element import NavigableString, PageElement, Tag
from Chrome import uc, get_driver
from Settings import SCROLL_PAUSE_TIME, MAX_SCROLL_ATTEMPTS, text_delimeter, MAX_OPEN_ATTEMPTS, output_file, \
    compare_start_ignore, accepted_length_diff, output_file_type, output_file_directory, \
    text_preview_symbols_number, parser_dict_default as PaDD, process_dict_default as PrDD, \
    sleeping_time, stylise_text, tesseract_path, link_info_part, jum_list, max_page_load_time, page_load_check_intervals
from Custom_exceptions import *
import requests
import time
from io import TextIOWrapper
import cssutils
from io import BytesIO
import re
import sys
from PIL import Image
import pytesseract
from Text_functions import string_with_meaning, convert_utf8_symbols, ask_for_authentification, dict_to_text, string_with_style
from Binary_converter import convert_binary
import validators

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from selenium.webdriver.common.by import By
from selenium.webdriver.remote.webdriver import WebElement

class FieldTypes(Enum):
    Text = 1
    Image = 2

class File():
    file_wrapper: object
    file_destination: str
    file_name: str
    file_type: str

    def create_full_path(self):
        return self.file_destination + self.file_name + '.' + self.file_type

    def __init__(self, path: str, name: str, type: str):
        self.file_destination = path
        self.file_name = name
        self.file_type = type
        self.file_wrapper = self.create_instance()

    def create_instance(self):
        if self.file_type == "txt":
            return open(self.create_full_path(), 'a', encoding='utf-8')
        elif self.file_type == "html":
            file = open(self.create_full_path(), 'w', encoding='utf-8')
            file.write("<html>\n<head></head>\n<body>\n")
            return file
        elif self.file_type == "docx":
            return Document()

    def write_text(self, string: str):
        if self.file_type == "docx":
            self.file_wrapper.add_paragraph(string)
        elif self.file_type == "txt":
            self.file_wrapper.write(string + "\n")
        elif self.file_type == "html":
            self.file_wrapper.write("<p>" + string + "</p>" + "\n")

    def write_image(self, data):
        if self.file_type == "docx":
            try:
                new_file = BytesIO()
                Image.open(BytesIO(convert_binary(data, "PIL"))).convert('RGB').save(new_file, format="png")

                p = self.file_wrapper.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                r = p.add_run()
                r.add_picture(new_file)#, width=Cm(image_width))
            except:
                return

    def close(self):
        if self.file_type == "docx":
            self.file_wrapper.save(self.create_full_path())
        elif self.file_type == "txt":
            self.file_wrapper.close()
        elif self.file_type == "html":
            self.file_wrapper.write("</body>")
            self.file_wrapper.close()



class Process:
    input_dict: dict

    chrome_used: bool
    scroll_used: bool
    soup: BeautifulSoup
    minimum_paragraphs: int
    minimum_length: int
    clear_cookies: bool
    driver: uc.Chrome
    open_attempts: int
    write_file: File
    to_sleep: bool

    block_screen: bool
    block_input_type: str
    block_input_limit: dict
    block_input_content: str
    block_button_type: str
    block_button_limit: dict


    def __init__(self, input: dict):
        self.input_dict = input
        self.open_attempts = 0

        self.chrome_used = input.get("chrome", PrDD.get("chrome"))
        self.scroll_used = input.get("scroll", PrDD.get("scroll"))
        self.minimum_paragraphs = input.get("min_par", PrDD.get("min_par"))
        self.minimum_length = input.get("min_len", PrDD.get("min_len"))
        self.clear_cookies = input.get("clearing", PrDD.get("clearing"))
        self.to_sleep = input.get("sleep", PrDD.get("sleep"))

        self.block_screen = input.get("block_screen", PrDD.get("block_screen"));
        self.block_input_type = input.get("input_type", PrDD.get("input_type"));
        self.block_input_limit = input.get("input_limit", PrDD.get("input_limit"));
        self.block_input_content = input.get("input_content", PrDD.get("input_content"));
        self.block_button_type = input.get("button_type", PrDD.get("button_type"));
        self.block_button_limit = input.get("button_limit", PrDD.get("button_limit"));

        self.initialise()

    def initialise(self):
        self.initialise_file()
        self.initialise_driver()

    def initialise_file(self):
        self.write_file = File(output_file_directory, output_file, output_file_type)

    def initialise_driver(self):
        self.driver = get_driver(self.input_dict) if self.chrome_used else None

    def get_soup(self, url: str):
        if self.to_sleep and self.open_attempts == 0:
            time.sleep(sleeping_time)

        self.open_attempts = 1
        if self.chrome_used:
            if "javascript" in url:
                self.driver.execute_script(url)
            elif self.driver.current_url == url:
                self.driver.refresh()
            else:
                self.driver.get(url)
            if self.scroll_used:
                self.execute_scrolling()

        if self.chrome_used:
            r = self.driver.page_source
        else:
            r = requests.get(url).content.decode('utf8')

        self.soup = BeautifulSoup(r, 'html.parser')
        return self.soup

    def execute_scrolling(self):
        attempts = 0
        last_height = self.driver.execute_script("return document.body.scrollHeight")
        while True:
            self.driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(SCROLL_PAUSE_TIME)
            new_height = self.driver.execute_script("return document.body.scrollHeight")
            attempts += 1
            if new_height == last_height or attempts > MAX_SCROLL_ATTEMPTS:
                break
            last_height = new_height

    def empty_cash(self):
        if self.driver != None:
            self.driver.delete_all_cookies()

    def closing(self):
        self.write_file.close()
        if self.chrome_used:
            if self.clear_cookies:
                self.empty_cash()
            self.driver.quit()

    def block_screen_talk(self):
        if self.block_input_type != None:
            input_field = self.driver.find_element(By.CSS_SELECTOR,
                f"{self.block_input_type}[{dict_to_text(self.block_input_limit, '=', ', ')}]")
            if input_field != None:
                input_field.send_keys(self.block_input_content)
            else:
                raise TargetNotFoundException("Failed to find input field")
        button_field = self.driver.find_element(By.CSS_SELECTOR,
                f"{self.block_button_type}[{dict_to_text(self.block_button_limit, '=', ', ')}]")
        if button_field != None:
            button_field.click()
        else:
            raise TargetNotFoundException("Failed to find button field")


    def refresh(self, url: str):
        if self.block_screen and self.chrome_used:
            try:
                self.block_screen_talk()
            except TargetNotFoundException:
                pass

            self.soup = self.get_soup(url)
            return self.soup
        else:
            current_attempts_num = self.open_attempts + 1
            self.empty_cash()
            if current_attempts_num > MAX_OPEN_ATTEMPTS:
                return None
            if current_attempts_num > 2 and self.chrome_used:
                self.driver.close()
                self.driver = get_driver(self.input_dict)

            self.soup = self.get_soup(url)
            self.open_attempts = current_attempts_num
            return self.soup
        '''
        if url.split('/')[2] == "tl.rulate.ru":
            if self.chrome_used:
                submit = self.driver.find_element(By.NAME, "ok")
                submit.click()
                self.soup = self.get_soup(url)
                return self.soup
            else:
                return None
        '''

    def write_text(self, string):
        if type(string) == str:
            string = convert_utf8_symbols(string)
            if stylise_text:
                string = string_with_style(string)
            self.write_file.write_text(string)
        elif type(string) == list:
            for part in string:
                self.write_text(part)
            #self.write_file.write_text(convert_utf8_symbols('\n'.join(string)) + '\n')
        elif type(string) == NavigableString:
            self.write_text(str(string.string))

    def write_image(self, data):
        if data is None:
            return
        if type(data) == list:
            for image in data:
                self.write_file.write_image(image)
        else:
            self.write_file.write_image(data)

    def update_content(self):
        if self.chrome_used:
            r = self.driver.page_source

        self.soup = BeautifulSoup(r, 'html.parser')
        return self.soup


class Parser:
    process: Process
    current_url: str
    soup: BeautifulSoup
    css_classes: dict
    page_open_time: float

    # extra parsing settings
    use_images: bool
    left_exclude_image: int
    right_exclude_image: int

    split_by_tags: list
    delimiter: str
    pre_phrase: str
    left_exclude: int
    right_exclude: int
    clean_empty: bool
    clean_equal: bool
    text_intelligent: bool
    text_expected_languages: list
    put_intelligent: bool

    text: list # used
    processed_text: list # used
    title: str # not used
    link: str # used
    images: list # used

    # parsing settings - search
    text_limit: dict
    title_limit: dict
    link_limit: dict

    text_holder: str
    title_holder: str
    link_holder: str

    title_pointer: int
    link_pointer: int

    text_cont: str
    title_cont: str
    link_cont: str

    #text_reader: type(easyocr.Reader)

    def __init__(self, process: Process, input: dict):
        self.process = process
        self.split_by_tags = input.get("tags_used", PaDD.get("tags_used"))
        self.use_images = input.get("images", PaDD.get("images"))
        self.left_exclude_image = input.get("left_image", PaDD.get("left_image"))
        self.right_exclude_image = input.get("right_image", PaDD.get("right_image"))

        self.left_exclude = input.get("left", PaDD.get("left"))
        self.right_exclude = input.get("right", PaDD.get("right"))
        self.delimiter = text_delimeter
        self.pre_phrase = ""
        self.current_url = input.get("url", "")
        self.clean_equal = input.get("clean_equal", PaDD.get("clean_equal"))
        self.clean_empty = input.get("clean_empty", PaDD.get("clean_empty"))

        self.text_holder = input.get("text_h", PaDD.get("text_h"))
        self.title_holder = input.get("title_h", PaDD.get("title_h"))
        self.link_holder = input.get("link_h", PaDD.get("link_h"))

        self.text_limit = input.get("text_l", PaDD.get("text_l"))
        self.title_limit = input.get("title_l", PaDD.get("title_l"))
        self.link_limit = input.get("link_l", PaDD.get("link_l"))

        self.title_pointer = input.get("title_p", PaDD.get("title_p"))
        self.link_pointer = input.get("link_p", PaDD.get("link_p"))

        self.text_cont = input.get("text_container", PaDD.get("text_container"))
        self.title_cont = input.get("title_container", PaDD.get("title_container"))
        self.link_cont = input.get("link_container", PaDD.get("link_container"))

        self.text_intelligent = input.get("text_intelligent", PaDD.get("text_intelligent"))
        self.text_expected_languages = input.get("text_lang", PaDD.get("text_lang"))
        self.put_intelligent = input.get("put_intelligent", PaDD.get("put_intelligent"))

        self.text_reader = None
        pytesseract.pytesseract.tesseract_cmd = tesseract_path

    def set_reader(self):
        if self.text_intelligent and self.text_expected_languages is not None:
            pass
            #self.text_reader = easyocr.Reader(self.text_expected_languages, gpu=True)

    def set_url(self, url: str, domain: str):
        if domain in url or "javascript" in url:
            self.current_url = url
        else:
            self.current_url = "https://" + domain + ('' if url.startswith('/') else '/') + url

    def process_page(self):
        try:
            self.process.open_attempts = 0
            self.soup = self.process.get_soup(self.current_url)
            self.page_open_time = time.time()
        except:
            raise Exception("Error with page opening")

        # opening page, waiting for it to load, refreshing if needed, checking if loaded successfully
        while True:
            try:
                self.retrieve_content()
                if self.check_pass() or ask_for_authentification("Content error, text: " +
                                                                 "\n".join(self.convert_to_strings(self.text))[:(text_preview_symbols_number
                                                                 if len("".join(self.convert_to_strings(self.text))) > text_preview_symbols_number
                                                                 else -1)] + "..."):
                    break
                else:
                    raise CommandException("Retry command")
            except HolderNotFoundException:
                if max_page_load_time == 0 or self.process.chrome_used is False:
                    raise Exception("Unsupported error")
                if time.time() - self.page_open_time > max_page_load_time:
                    raise MaxOpeningTimeExceeded("Maximum opening time was exceeded while opening " + self.current_url)
                time.sleep(page_load_check_intervals)
                self.soup = self.process.update_content()
            except CommandException:
                self.soup = self.process.refresh(self.current_url)
                self.page_open_time = time.time()
                if self.soup is None:
                    raise Exception("Unsupported error, maybe error with overcoming defense")

        # remake for ordering by line saved in soup element
        # page is known to be loaded successfully
        self.get_true_content()
        # get title and write it
        self.process.write_text(self.retrieve_title())
        # get images if they are needed
        if self.use_images:
            self.retrieve_images()
        # write info in the file
        if not self.put_intelligent or not self.use_images:
            self.text = self.convert_to_strings(self.text)
            self.use_text()
            if self.use_images:
                self.get_images_content()
                self.process.write_image(self.images)
        else:
            self.combined_write()
        # get and use link
        return self.retrieve_link()

    def combined_write(self):
        # combine self.text and self.images and write them in the file in the suitable order
        ordered_content = []
        self.processed_text = self.text
        for textUnit in self.text:
            if isinstance(textUnit, Tag):
                ordered_content.append([textUnit.sourceline, textUnit.sourcepos, FieldTypes.Text, textUnit.text])
            elif isinstance(textUnit, str):
                located = False
                for index in range(len(ordered_content), -1, -1):
                    if ordered_content[index][2] == FieldTypes.Text:
                        ordered_content[index][3] += "\n" + textUnit
                        located = True
                        break
                if not located:
                    maximumLine = max(ordered_content, lambda x: x[0]) if len(ordered_content) > 0 else 10000000000
                    maximumPos = max(ordered_content, lambda x: x[1]) if len(ordered_content) > 0 else 10000000000
                    ordered_content.append([maximumLine + 1, maximumPos + 1, FieldTypes.Text, textUnit])

        for imageUnit in self.images:
            if isinstance(imageUnit, Tag):
                ordered_content.append([imageUnit.sourceline, imageUnit.sourcepos, FieldTypes.Image, self.get_image_content(imageUnit)])

        ordered_content.sort(key=lambda x: (x[0], x[1]))

        for unit in ordered_content:
            if unit[2] == FieldTypes.Text:
                text = unit[3]
                answer = self.process_text_unit(text)
                if answer is not None:
                    self.process.write_text(answer)
                else:
                    self.processed_text.pop(0) # just to show how many were erased
            elif unit[2] == FieldTypes.Image:
                image = unit[3]
                self.process.write_image(image)

    def retrieve_images(self):
        if self.text_limit != None:
            self.images = self.get_search_results(holder_type = 'div', holder_limit = self.text_limit, target_type = "img", accept_zero = True)
        else:
            self.images = self.get_search_results(target_type = "img", accept_zero = True)
        self.get_true_content_image()
        #self.get_images_content()

    def get_image_content(self, image: PageElement):
        try:
            for param in link_info_part:
                info = image.get(param)
                if info is None:
                    continue
                if validators.url(info):
                    return requests.get(info).content
                info = "/".join(self.current_url.split('/')[:3]) + info
                if validators.url(info):
                    return requests.get(info).content
        except:
            return None

    def get_images_content(self):
        for index in range(0, len(self.images)):
            self.images[index] = self.get_image_content(self.images[index])
        self.images = list(self.images)

    def retrieve_content(self):  # soup already received
        if self.split_by_tags != None:
            unit = self.get_search_results(target_type = 'div', target_limit = self.text_limit, target_pointer = 0)
            index = 0
            while index < len(unit.contents):
                if type(unit.contents[index]) == Tag:
                    if unit.contents[index].name in self.split_by_tags and not string_with_meaning(unit.contents[index].text):
                        unit.contents[index].replaceWith(self.delimiter)
                        index+=1
                    else:
                        unit.contents[index].unwrap()
                else:
                    index+=1
            self.text = re.sub(f"[{self.delimiter}]+", self.delimiter, unit.get_text()).split(self.delimiter)
        else:
            if self.text_limit != None:
                if self.text_cont is not None:
                    self.text = self.get_search_results(holder_type = self.text_cont, holder_limit = self.text_limit, target_type = self.text_holder)
                else:
                    self.text = self.get_search_results(target_limit = self.text_limit, target_type = self.text_holder)
            else:
                self.text = self.get_search_results(target_type = self.text_holder)
        if self.text_intelligent:
            self.parse_styles()
            self.intelligent_filter(self.text)
        # self.text = self.convert_to_strings(self.text)

    def convert_to_strings(self, string_list: list):
        try:
            result = []
            for x in string_list:
                if type(x) == str:
                    result.append(x)
                elif type(x) == NavigableString:
                    result.append(x.text)
                elif type(x) == Tag:
                    x: Tag
                    writer = io.StringIO()
                    for elem in x.contents:
                        if type(elem) == NavigableString:
                            writer.write(elem.text)
                        elif type(elem) == Tag and elem.name in ['span', 'em', 'strong']:
                            writer.write(elem.text)
                        elif type(elem) == Tag and elem.name == 'br':
                            writer.write("\n")
                    result.append(writer.getvalue())
            return result
        except:
            self.process.closing()
            raise Exception("Unexpected text container")

    def check_pass(self):
        return (len(self.text) >= self.process.minimum_paragraphs or
                sum(sys.getsizeof(x) for x in self.text) > self.process.minimum_length)

    def retrieve_title(self):
        try:
            if self.title_holder == "title":
                return self.soup.title.string
            if self.title_holder == "empty":
                return ""
            if self.title_limit != None and self.title_cont != None:
                return self.get_search_results(holder_type = self.title_cont, holder_limit = self.title_limit, target_type = self.title_holder, target_pointer = self.title_pointer).text
            elif self.title_limit != None:  # limit by it itself
                return self.get_search_results(target_type = self.title_holder, target_limit = self.title_limit, target_pointer = self.title_pointer).text
            elif self.title_cont != None:
                return self.get_search_results(holder_type = self.title_cont, target_type=self.title_holder,
                                               target_pointer=self.title_pointer).text
            else:
                return self.get_search_results(target_type=self.title_holder, target_pointer=self.title_pointer).text
        except:
            return ""  # to remove
            self.process.closing()
            raise Exception("Title was not detected")

    def retrieve_link(self):
        if self.link_limit != None:
            if self.link_cont != None:
                return self.get_search_results(holder_type = self.link_cont, holder_limit = self.link_limit, target_type = self.link_holder, target_pointer = self.link_pointer).get('href')
            else:
                return self.get_search_results(target_type = self.link_holder, target_limit = self.link_limit, target_pointer = self.link_pointer).get('href')
        else:
            return self.get_search_results(target_type = self.link_holder, target_pointer = self.link_pointer).get('href')

    def get_true_content(self):
        if (len(self.text) > self.left_exclude and self.left_exclude != 0):
            self.text = self.text[self.left_exclude:]
        if (len(self.text) > self.right_exclude and self.right_exclude != 0):
            self.text = self.text[:self.right_exclude * -1]

    def get_true_content_image(self):
        if (len(self.images) > self.left_exclude_image and self.left_exclude_image != 0):
            self.images = self.images[self.left_exclude_image:]
        if (len(self.images) > self.right_exclude_image and self.right_exclude_image != 0):
            self.images = self.images[:self.right_exclude_image * -1]

    def process_text_unit(self, text: str):
        if text == None or text == "":
            None
        processed_text = text
        processed_text = processed_text.replace("\n", ' ')
        if self.clean_empty and not string_with_meaning(processed_text):
            return None
        if self.clean_equal and not self.repeat_check(processed_text):
            return None
        return processed_text

    def use_text(self):
        if self.text == None or self.text == "":
            return
        self.processed_text = self.text
        self.processed_text = [line.replace("\n", ' ') for line in self.processed_text]
        if self.clean_empty:
            self.processed_text = [line for line in self.processed_text if string_with_meaning(line)]
        if self.clean_equal:
            self.processed_text = [line for line in self.processed_text if self.repeat_check(line)]
        # now write to file
        self.process.write_text(self.processed_text)

    def repeat_check(self, text: str):  # shows if text was repeated
        if abs(len(self.pre_phrase) - len(text)) > accepted_length_diff or \
                self.pre_phrase.find(text[compare_start_ignore:]) == -1 or \
                text.find(self.pre_phrase[compare_start_ignore:]) == -1:
            self.update_pre_phrase(text)
            return True
        else:
            print(f"Repetition of {text} founded - text was ignored")
            self.update_pre_phrase(text)
            return False

    def update_pre_phrase(self, new_phrase: str):  # updates previous phrase based on empty and settings
        if self.clean_empty:
            # empty string were emptied
            self.pre_phrase = new_phrase
            return
        else:
            # empty string is expected
            if string_with_meaning(new_phrase):
                self.pre_phrase = new_phrase
            return

    def get_search_results(self, holder_type = None, holder_limit = None, target_type = None, target_limit = None, target_pointer = None, accept_zero = False):
        if target_type == None:
            raise UnsupportedArgumentsException("Search type wasn't given")
        holder = None
        if holder_type != None:
            if holder_limit != None:
                holder = self.soup.find(holder_type, holder_limit)
                if holder == None:
                    holder = self.soup.find(holder_type, **holder_limit)
            else:
                holder = self.soup.find(holder_type)
            if holder == None:
                raise HolderNotFoundException(f"Unable to find holder using {holder_type} and {holder_limit}")
        else:
            holder = self.soup

        results = None
        if target_limit != None:
            results = holder.find_all(target_type, target_limit)
            if (target_pointer != None and len(results) <= target_pointer) or (len(results) == 0):  # error check
                results = holder.find_all(target_type, **target_limit)
        else:
            results = holder.find_all(target_type)
        if (target_pointer != None and len(results) <= target_pointer) or ((len(results) == 0) and not accept_zero):  # error check
            raise TargetNotFoundException(f"Unable to find target using holder {holder_type} and {holder_limit}, target {target_type} - {target_limit} - {target_pointer}")
        if target_pointer != None:
            return results[target_pointer]
        else:
            return results

    def intelligent_filter(self, elements: list):
        for index in range(len(elements) - 1, -1, -1):
            element: Tag = elements[index]
            if not self.check_visibility(element):
                elements.pop(index)
                continue
            if self.process.chrome_used and self.jummed_text(element):
                elements[index] = self.get_view_content(element)
                continue

            for child in element.contents:
                if type(child) == NavigableString:
                    continue
                if not self.check_visibility(child):
                    child.extract()
                    continue
                if self.process.chrome_used and self.jummed_text(child):
                    child.replace_with(self.get_view_content(child))
                    #child.text = self.get_view_content(child)

    def check_visibility(self, element: PageElement):
        style_name = "style"
        class_name = "class"
        attributes = dict(element.attrs)
        parsed_properties = {}
        if style_name in attributes.keys():
            s = cssutils.parseStyle(attributes[style_name])
            attributes.pop(style_name)
            attributes.update(s)
        if class_name in attributes.keys():
            for css_class in attributes[class_name]:
                if ('.' + css_class) not in self.css_classes: continue
                for k, v in self.css_classes['.' + css_class].items():
                    parsed_properties[k] = v
        return not (attributes.get("overflow", "Yes") == "hidden" or
                    element.get("display", "Yes") == "none" or
                    parsed_properties.get("display", "Yes") == "none")

    def get_view_content(self, element: Tag):
        #if (self.text_reader == None):
        #    self.set_reader()
        xpath = self.xpath_soup(element)
        self.unwrap_xpath(xpath)
        obj = self.process.driver.find_element(By.XPATH, xpath)
        #self.process.driver.execute_script("arguments[0].setAttribute(arguments[1], arguments[2]);", obj, style_for_padding.get("type"), style_for_padding.get("value")) # почему-то делает отступ у первой строки и левая граница берется по ней
        self.process.driver.execute_script("arguments[0].scrollIntoView(true);", obj)

        #image_data = base64.b64decode(obj.screenshot_as_base64)
        image_data = self.get_screen_of_element(obj, 5, 5)

        result = pytesseract.image_to_string(Image.open(BytesIO(image_data)), lang=self.text_expected_languages)
        #result = self.text_reader.readtext(image_data, paragraph = True, detail = 0, mag_ratio = 2, blocklist = [';', ':', '|'], low_text = 0.25)
        return re.sub(r"\n+", " ", result).replace('|', 'I')

    def get_screen_of_element(self, element: WebElement, offset_x: int, offset_y: int):
        normalise = lambda val, hor, limits: min(max(val, limits[0][0] if hor else limits[0][1]), limits[1][0] if hor else limits[1][1])
        l = element.location_once_scrolled_into_view
        s = element.size
        limits = []
        limits.append([0, 0])
        limits.append([self.process.driver.get_window_size().get('width'), self.process.driver.get_window_size().get('height')])
        tup = (
            normalise(l['x'] - offset_x, True, limits),
            normalise(l['y'] - offset_y, False, limits),
            normalise(l['x'] + s['width'] + offset_x * 2, True, limits),
            normalise(l['y'] + s['height'] + offset_y * 2, False, limits)
        )

        screen_image = base64.b64decode(self.process.driver.get_screenshot_as_base64())
        image = Image.open(io.BytesIO(screen_image))
        scale = image.size[0] / limits[1][0] # shows scale of image relative to screen by width (should be the same by height)
        tup = tuple((int)(scale * elem) for elem in tup)
        image = image.crop(tup)
        # image.show()  # comment later
        buffer = io.BytesIO()
        image.save(buffer, format="PNG")
        return buffer.getvalue()


    def jummed_text(self, element: Tag):
        return (element.has_attr("class") and any(jum_class in element["class"] for jum_class in jum_list)) or any((type(elem) == Tag and self.jummed_text(elem)) for elem in element.contents)

    def xpath_soup(self, element):  # from https://gist.github.com/ergoithz/6cf043e3fdedd1b94fcf
        # type: (typing.Union[bs4.element.Tag, bs4.element.NavigableString]) -> str
        components = []
        child = element if element.name else element.parent
        for parent in child.parents:  # type: bs4.element.Tag
            siblings = parent.find_all(child.name, recursive=False)
            components.append(
                child.name if 1 == len(siblings) else '%s[%d]' % (
                    child.name,
                    next(i for i, s in enumerate(siblings, 1) if s is child)
                )
            )
            child = parent
        components.reverse()
        return '/%s' % '/'.join(components)

    def unwrap_xpath(self, xpath):
        parts = xpath.split('/')
        current_path = parts[0]
        for part in parts[1:]:
            current_path += '/' + part
            # unwrap details
            if part.count('details') == 1:
                obj = self.process.driver.find_element(By.XPATH, current_path)
                self.process.driver.execute_script("arguments[0].open = true;", obj)

    def parse_styles(self):
        self.css_classes = {}
        for styles in self.soup.select('style'):
            css = cssutils.parseString(styles.encode_contents(), validate=False)
            for rule in css:
                if rule.type == rule.STYLE_RULE:
                    style = rule.selectorText
                    self.css_classes[style] = {}
                    for item in rule.style:
                        propertyname = item.name
                        value = item.value
                        self.css_classes[style][propertyname] = value

